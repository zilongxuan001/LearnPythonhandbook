#调整word图片大小

# 导入模块
from docx import Document
# 此模块中包含 docx 中各类单位方法
from docx import shared

doc = Document()

# 在文档中增加图片,并对设置图片大小
# 当只设置一个方向的长度（宽或高）时，另一方向会自动缩放

doc.add_picture('1.png',height=shared.Cm(11.25))  # 按厘米设置
doc.add_picture('1.png',width=shared.Cm(18.04)) 

# 保存文件
doc.save('test2.docx')